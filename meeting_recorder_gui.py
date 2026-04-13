#!/usr/bin/env python3
"""
Meeting Recorder — Interfaccia Grafica
Sviluppato da Marco Bonometti
"""

import os
import sys
import wave
import json
import threading
import datetime
import subprocess
from pathlib import Path

# ── Fix PyInstaller multiprocessing ──────────────────
if getattr(sys, 'frozen', False):
    import multiprocessing
    multiprocessing.freeze_support()

# ── CRITICO: importa ctranslate2/faster_whisper PRIMA di PyQt6
# Su Windows c'è un conflitto tra le DLL di PyQt6 e ctranslate2
# se ctranslate2 viene caricato dopo PyQt6
try:
    import ctranslate2
    from faster_whisper import WhisperModel as _WhisperModel
except Exception:
    pass

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFileDialog, QTextEdit, QLineEdit,
    QTabWidget, QFrame, QScrollArea, QSizePolicy, QMessageBox,
    QProgressBar, QListWidget, QListWidgetItem, QSplitter,
    QGroupBox, QFormLayout, QSpinBox, QComboBox, QStatusBar, QDialog,
    QCheckBox,
)
from PyQt6.QtCore import (
    Qt, QThread, pyqtSignal, QTimer, QSize, QPropertyAnimation,
    QEasingCurve,
)
from PyQt6.QtGui import (
    QFont, QColor, QPalette, QIcon, QPixmap, QPainter,
    QLinearGradient, QBrush, QPen, QFontDatabase,
)

# ──────────────────────────────────────────────────────
# CONFIGURAZIONE DEFAULT
# ──────────────────────────────────────────────────────
CONFIG_FILE     = Path.home() / ".meeting_recorder_config.json"
OUTPUT_DIR      = Path.home() / "Documenti" / "Riunioni"
LM_STUDIO_URL   = "http://localhost:1234/v1/chat/completions"
LM_STUDIO_MODEL = "local-model"
WHISPER_MODEL   = "medium"
LANGUAGE        = "it"
SAMPLE_RATE     = 16000
CHANNELS        = 1
CHUNK           = 1024

# Debug log
_DEBUG_LOG = Path.home() / "meeting_recorder_debug.txt"
def _dbg(msg):
    try:
        with open(_DEBUG_LOG, "a", encoding="utf-8") as f:
            f.write(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {msg}\n")
    except Exception:
        pass

def _format_ts(seconds):
    """Formatta secondi in MM:SS o HH:MM:SS."""
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h > 0:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"

# ──────────────────────────────────────────────────────
# STILE DARK PROFESSIONALE
# ──────────────────────────────────────────────────────
STYLE = """
QMainWindow, QWidget {
    background-color: #0F1117;
    color: #E2E8F0;
    font-family: 'SF Pro Display', 'Segoe UI', Arial, sans-serif;
}
QTabWidget::pane {
    border: 1px solid #2D3748;
    background-color: #161B27;
    border-radius: 8px;
}
QTabBar::tab {
    background: #1A2035;
    color: #94A3B8;
    padding: 10px 24px;
    border: none;
    font-size: 13px;
    font-weight: 500;
    border-radius: 6px 6px 0 0;
    margin-right: 2px;
}
QTabBar::tab:selected {
    background: #2563EB;
    color: #FFFFFF;
    font-weight: 600;
}
QTabBar::tab:hover:!selected {
    background: #1E3A5F;
    color: #CBD5E1;
}
QPushButton {
    background-color: #2563EB;
    color: white;
    border: none;
    border-radius: 8px;
    padding: 10px 20px;
    font-size: 13px;
    font-weight: 600;
}
QPushButton:hover { background-color: #1D4ED8; }
QPushButton:pressed { background-color: #1E40AF; }
QPushButton:disabled { background-color: #374151; color: #6B7280; }
QPushButton#btnDanger {
    background-color: #DC2626;
}
QPushButton#btnDanger:hover { background-color: #B91C1C; }
QPushButton#btnSuccess {
    background-color: #059669;
}
QPushButton#btnSuccess:hover { background-color: #047857; }
QPushButton#btnSecondary {
    background-color: #1E293B;
    border: 1px solid #334155;
    color: #CBD5E1;
}
QPushButton#btnSecondary:hover { background-color: #263548; }
QLineEdit, QComboBox, QSpinBox {
    background-color: #1E293B;
    border: 1px solid #334155;
    border-radius: 6px;
    padding: 8px 12px;
    color: #E2E8F0;
    font-size: 13px;
}
QLineEdit:focus, QComboBox:focus, QSpinBox:focus {
    border: 1px solid #2563EB;
}
QComboBox::drop-down { border: none; width: 24px; }
QComboBox::down-arrow { width: 12px; height: 12px; }
QComboBox QAbstractItemView {
    background-color: #1E293B;
    border: 1px solid #334155;
    color: #E2E8F0;
    selection-background-color: #2563EB;
}
QTextEdit {
    background-color: #0D1117;
    border: 1px solid #21262D;
    border-radius: 8px;
    color: #C9D1D9;
    font-family: 'Menlo', 'Consolas', 'Courier New', monospace;
    font-size: 12px;
    padding: 12px;
    line-height: 1.6;
}
QListWidget {
    background-color: #161B27;
    border: 1px solid #2D3748;
    border-radius: 8px;
    color: #E2E8F0;
    font-size: 13px;
}
QListWidget::item {
    padding: 10px 14px;
    border-bottom: 1px solid #1E293B;
    border-radius: 4px;
}
QListWidget::item:selected {
    background-color: #1E3A5F;
    color: #60A5FA;
}
QListWidget::item:hover:!selected {
    background-color: #1A2540;
}
QScrollBar:vertical {
    background: #161B27;
    width: 8px;
    border-radius: 4px;
}
QScrollBar::handle:vertical {
    background: #334155;
    border-radius: 4px;
    min-height: 30px;
}
QScrollBar::handle:vertical:hover { background: #475569; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; }
QGroupBox {
    border: 1px solid #2D3748;
    border-radius: 8px;
    margin-top: 12px;
    padding-top: 8px;
    font-size: 13px;
    font-weight: 600;
    color: #94A3B8;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 12px;
    padding: 0 6px;
    color: #60A5FA;
}
QProgressBar {
    background-color: #1E293B;
    border: none;
    border-radius: 4px;
    height: 6px;
    text-align: center;
}
QProgressBar::chunk {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 #2563EB, stop:1 #7C3AED);
    border-radius: 4px;
}
QStatusBar {
    background-color: #0D1117;
    color: #64748B;
    border-top: 1px solid #1E293B;
    font-size: 12px;
}
QSplitter::handle { background-color: #1E293B; width: 2px; }
QFrame#card {
    background-color: #161B27;
    border: 1px solid #1E293B;
    border-radius: 12px;
}
QLabel#title {
    color: #F1F5F9;
    font-size: 22px;
    font-weight: 700;
}
QLabel#subtitle {
    color: #64748B;
    font-size: 12px;
}
QLabel#stat_value {
    color: #60A5FA;
    font-size: 28px;
    font-weight: 700;
}
QLabel#stat_label {
    color: #64748B;
    font-size: 11px;
    font-weight: 500;
}
QLabel#recording {
    color: #EF4444;
    font-size: 14px;
    font-weight: 700;
}
QLabel#success {
    color: #10B981;
    font-size: 13px;
}
QLabel#warning {
    color: #F59E0B;
    font-size: 13px;
}
"""

# ──────────────────────────────────────────────────────
# CONFIG MANAGER
# ──────────────────────────────────────────────────────

# Cache globale del modello Whisper — caricato una volta sola
_WHISPER_MODEL_CACHE = {}

def get_whisper_model(model_name):
    """Carica il modello Whisper una volta e lo riusa."""
    global _WHISPER_MODEL_CACHE
    if model_name not in _WHISPER_MODEL_CACHE:
        _dbg(f"Caricamento modello Whisper: {model_name}")
        from faster_whisper import WhisperModel
        cache_dir = str(Path.home() / ".cache" / "huggingface" / "hub")
        _WHISPER_MODEL_CACHE[model_name] = WhisperModel(
            model_name,
            device="cpu",
            compute_type="int8",
            download_root=cache_dir,
        )
        _dbg("Modello Whisper pronto")
    return _WHISPER_MODEL_CACHE[model_name]
def load_config():
    defaults = {
        "lm_url":        LM_STUDIO_URL,
        "lm_model":      LM_STUDIO_MODEL,
        "whisper_model": WHISPER_MODEL,
        "language":      LANGUAGE,
        "output_dir":    str(Path.home() / "Documents" / "Riunioni"),
    }
    if CONFIG_FILE.exists():
        try:
            data = json.loads(CONFIG_FILE.read_text())
            defaults.update(data)
        except Exception:
            pass
    return defaults

def save_config(cfg):
    try:
        CONFIG_FILE.write_text(json.dumps(cfg, indent=2))
    except Exception:
        pass

# ──────────────────────────────────────────────────────
# WORKER THREADS
# ──────────────────────────────────────────────────────
class RecordWorker(QThread):
    finished   = pyqtSignal(str)   # path al file wav
    error      = pyqtSignal(str)
    level      = pyqtSignal(float) # livello audio 0-1

    def __init__(self, output_dir, cfg):
        super().__init__()
        self.output_dir = Path(output_dir)
        self.cfg = cfg
        self._stop = False
        self.frames = []

    def stop(self):
        self._stop = True

    def run(self):
        try:
            import numpy as np
            import math
            import sounddevice as sd

            self.frames = []

            if sys.platform == "darwin":
                # ── macOS: sounddevice ──────────────────────────
                def callback(indata, frames, time, status):
                    try:
                        self.frames.append(indata.copy().tobytes())
                        data_f = indata.astype(np.float32)
                        mean_sq = float(np.mean(data_f ** 2))
                        rms = math.sqrt(mean_sq) if mean_sq > 0 else 0.0
                        self.level.emit(min(rms * 10, 1.0))
                    except Exception:
                        self.level.emit(0.0)

                with sd.InputStream(
                    samplerate=SAMPLE_RATE,
                    channels=1,
                    dtype='int16',
                    blocksize=CHUNK,
                    callback=callback,
                ):
                    while not self._stop:
                        sd.sleep(100)

            else:
                # ── Windows: sounddevice con MME ────────────────
                try:
                    devs = sd.query_devices()
                    hostapis = sd.query_hostapis()
                    mme_idx = next((i for i, h in enumerate(hostapis) if 'MME' in h['name']), 0)
                    device_id = next(
                        (i for i, d in enumerate(devs)
                         if d['max_input_channels'] > 0 and d['hostapi'] == mme_idx),
                        None
                    )
                except Exception:
                    device_id = None

                def callback(indata, frames, time, status):
                    try:
                        self.frames.append(indata.copy().tobytes())
                        data_f = indata.astype(np.float32)
                        mean_sq = float(np.mean(data_f ** 2))
                        rms = math.sqrt(mean_sq) if mean_sq > 0 else 0.0
                        self.level.emit(min(rms * 10, 1.0))
                    except Exception:
                        self.level.emit(0.0)

                stream_kwargs = dict(
                    samplerate=SAMPLE_RATE,
                    channels=1,
                    dtype='int16',
                    blocksize=CHUNK,
                    callback=callback,
                )
                if device_id is not None:
                    stream_kwargs['device'] = device_id

                with sd.InputStream(**stream_kwargs):
                    while not self._stop:
                        sd.sleep(100)
            # Salva il file WAV
            self.output_dir.mkdir(parents=True, exist_ok=True)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            wav_path = self.output_dir / f"riunione_{ts}.wav"
            with wave.open(str(wav_path), "wb") as wf:
                wf.setnchannels(CHANNELS)
                wf.setsampwidth(2)  # int16 = 2 bytes
                wf.setframerate(SAMPLE_RATE)
                wf.writeframes(b"".join(self.frames))
            self.finished.emit(str(wav_path))
        except Exception as e:
            self.error.emit(str(e))


class TranscribeWorker(QThread):
    progress   = pyqtSignal(str)
    finished   = pyqtSignal(str)
    error      = pyqtSignal(str)

    def __init__(self, audio_path, cfg):
        super().__init__()
        self.audio_path = audio_path
        self.cfg = cfg

    # ── Speaker diarization helpers ──────────────────────
    @staticmethod
    def _assign_speakers_by_pause(segments_list, pause_threshold=1.5):
        """Euristica: cambia speaker quando c'è una pausa > soglia tra segmenti.
           Restituisce lista di (speaker_label, text, start, end)."""
        if not segments_list:
            return []
        result = []
        current_speaker = 1
        prev_end = 0.0
        for seg in segments_list:
            gap = seg["start"] - prev_end
            if gap > pause_threshold and prev_end > 0:
                current_speaker += 1
            result.append((f"Speaker {current_speaker}", seg["text"], seg["start"], seg["end"]))
            prev_end = seg["end"]
        return result

    @staticmethod
    def _try_pyannote_diarization(audio_path):
        """Tenta diarizzazione con pyannote.audio. Ritorna Annotation o None."""
        try:
            from pyannote.audio import Pipeline
            import torch
            hf_token = os.environ.get("HF_TOKEN", "")
            if not hf_token:
                _dbg("pyannote: HF_TOKEN non impostato, skip diarizzazione avanzata")
                return None
            pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                token=hf_token,
            )
            device = "cuda" if torch.cuda.is_available() else "cpu"
            pipeline.to(torch.device(device))
            result = pipeline(audio_path)
            _dbg("pyannote diarization completata")
            # pyannote 3.x restituisce DiarizeOutput, estraiamo l'Annotation
            if hasattr(result, 'speaker_diarization'):
                return result.speaker_diarization
            # pyannote 2.x restituisce direttamente Annotation
            return result
        except ImportError:
            _dbg("pyannote.audio non installato — uso euristica pause")
            return None
        except Exception as e:
            _dbg(f"pyannote errore: {e} — fallback a euristica")
            return None

    @staticmethod
    def _merge_diarization_with_segments(diarization, segments_list):
        """Associa ogni segmento Whisper allo speaker pyannote più sovrapposto."""
        result = []
        for seg in segments_list:
            seg_start, seg_end = seg["start"], seg["end"]
            best_speaker = "Speaker ?"
            best_overlap = 0.0
            for turn, _, speaker in diarization.itertracks(yield_label=True):
                overlap_start = max(seg_start, turn.start)
                overlap_end = min(seg_end, turn.end)
                overlap = max(0, overlap_end - overlap_start)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_speaker = speaker
            result.append((best_speaker, seg["text"], seg_start, seg_end))
        return result

    @staticmethod
    def _format_diarized_text(diarized_segments):
        """Formatta i segmenti diarizzati in testo leggibile con timestamp."""
        lines = []
        current_speaker = None
        current_texts = []
        current_start = 0.0
        for speaker, text, start, end in diarized_segments:
            if speaker != current_speaker:
                if current_speaker is not None and current_texts:
                    ts = _format_ts(current_start)
                    lines.append(f"[{ts}] [{current_speaker}]: {' '.join(current_texts)}")
                current_speaker = speaker
                current_texts = [text.strip()]
                current_start = start
            else:
                current_texts.append(text.strip())
        if current_speaker is not None and current_texts:
            ts = _format_ts(current_start)
            lines.append(f"[{ts}] [{current_speaker}]: {' '.join(current_texts)}")
        return "\n\n".join(lines)

    def run(self):
        try:
            _dbg("TranscribeWorker.run START")
            self.progress.emit("Caricamento modello Whisper...")

            cache_dir = str(Path.home() / ".cache" / "huggingface" / "hub")
            model_name = self.cfg["whisper_model"]
            language   = self.cfg["language"]
            audio_path = self.audio_path
            diarize    = self.cfg.get("diarization", True)

            if sys.platform == "darwin":
                # ── Mac: trascrizione diretta nel thread ──
                _dbg(f"Mac: caricamento diretto modello {model_name}")
                from faster_whisper import WhisperModel
                model = WhisperModel(
                    model_name,
                    device="cpu",
                    compute_type="int8",
                    download_root=cache_dir,
                )
                _dbg("modello caricato OK")
                self.progress.emit(f"Trascrizione in corso (modello: {model_name})...")
                lang = None if language == "auto" else language
                segments, info = model.transcribe(audio_path, language=lang, vad_filter=True)
                segments_list = [{"text": seg.text, "start": seg.start, "end": seg.end} for seg in segments]

                if diarize and segments_list:
                    self.progress.emit("Identificazione speaker in corso...")
                    diarization = self._try_pyannote_diarization(audio_path)
                    if diarization is not None:
                        diarized = self._merge_diarization_with_segments(diarization, segments_list)
                    else:
                        diarized = self._assign_speakers_by_pause(segments_list)
                    text = self._format_diarized_text(diarized)
                else:
                    text = " ".join(s["text"].strip() for s in segments_list)

                word_count = len(text.split())
                _dbg(f"trascrizione OK: {word_count} parole, diarize={diarize}")
                self.progress.emit(f"Trascrizione completata — {word_count} parole")
                self.finished.emit(text)

            else:
                # ── Windows: subprocess con Python esterno ──
                import subprocess, os, shutil

                if getattr(sys, 'frozen', False):
                    python_exe = None
                    username = os.environ.get("USERNAME", "mark")
                    direct_paths = [
                        rf"C:\Users\{username}\AppData\Local\Programs\Python\Python311\python.exe",
                        rf"C:\Users\{username}\AppData\Local\Programs\Python\Python310\python.exe",
                        rf"C:\Users\{username}\AppData\Local\Programs\Python\Python312\python.exe",
                        r"C:\Python311\python.exe",
                        r"C:\Python310\python.exe",
                        r"C:\Program Files\Python311\python.exe",
                    ]
                    for p in direct_paths:
                        if os.path.exists(p):
                            python_exe = p
                            _dbg(f"Python trovato in: {p}")
                            break
                    if not python_exe:
                        for name in ["python3.11", "python3.10", "python"]:
                            found = shutil.which(name)
                            if found and "WindowsApps" not in found and "Microsoft" not in found:
                                python_exe = found
                                break
                    if not python_exe:
                        self.error.emit("Python non trovato.\nInstalla Python 3.11 da python.org e spunta 'Add Python to PATH'")
                        return
                else:
                    python_exe = sys.executable

                _dbg(f"Python exe: {python_exe}")
                script = f"""
import json, os, sys
from pathlib import Path
from faster_whisper import WhisperModel

cache_dir = {repr(cache_dir)}
model = WhisperModel({repr(model_name)}, device="cpu", compute_type="int8", download_root=cache_dir)
lang = None if {repr(language)} == "auto" else {repr(language)}
segments, info = model.transcribe({repr(audio_path)}, language=lang, vad_filter=True)
segments_list = [{{"text": seg.text, "start": seg.start, "end": seg.end}} for seg in segments]

diarize = {repr(diarize)}

def format_ts(seconds):
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h > 0:
        return f"{{h:02d}}:{{m:02d}}:{{s:02d}}"
    return f"{{m:02d}}:{{s:02d}}"

def assign_speakers_by_pause(segs, pause_threshold=1.5):
    if not segs:
        return []
    result = []
    current_speaker = 1
    prev_end = 0.0
    for seg in segs:
        gap = seg["start"] - prev_end
        if gap > pause_threshold and prev_end > 0:
            current_speaker += 1
        result.append((f"Speaker {{current_speaker}}", seg["text"], seg["start"], seg["end"]))
        prev_end = seg["end"]
    return result

if diarize and segments_list:
    # Prova pyannote
    diarization = None
    try:
        hf_token = os.environ.get("HF_TOKEN", "")
        if hf_token:
            from pyannote.audio import Pipeline
            import torch
            pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                token=hf_token,
            )
            device = "cuda" if torch.cuda.is_available() else "cpu"
            pipeline.to(torch.device(device))
            diarization = pipeline({repr(audio_path)})
            # pyannote 3.x: estrai Annotation da DiarizeOutput
            if hasattr(diarization, 'speaker_diarization'):
                diarization = diarization.speaker_diarization
    except Exception:
        pass

    if diarization is not None:
        diarized = []
        for seg in segments_list:
            seg_start, seg_end = seg["start"], seg["end"]
            best_speaker = "Speaker ?"
            best_overlap = 0.0
            for turn, _, speaker in diarization.itertracks(yield_label=True):
                overlap_start = max(seg_start, turn.start)
                overlap_end = min(seg_end, turn.end)
                overlap = max(0, overlap_end - overlap_start)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_speaker = speaker
            diarized.append((best_speaker, seg["text"], seg_start, seg_end))
    else:
        diarized = assign_speakers_by_pause(segments_list)

    # Formatta output con raggruppamento per speaker
    lines = []
    current_speaker = None
    current_texts = []
    current_start = 0.0
    for speaker, text, start, end in diarized:
        if speaker != current_speaker:
            if current_speaker is not None and current_texts:
                ts = format_ts(current_start)
                lines.append(f"[{{ts}}] [{{current_speaker}}]: {{' '.join(current_texts)}}")
            current_speaker = speaker
            current_texts = [text.strip()]
            current_start = start
        else:
            current_texts.append(text.strip())
    if current_speaker is not None and current_texts:
        ts = format_ts(current_start)
        lines.append(f"[{{ts}}] [{{current_speaker}}]: {{' '.join(current_texts)}}")
    print("\\n\\n".join(lines))
else:
    text = " ".join(s["text"].strip() for s in segments_list)
    print(text)
"""
                _dbg("Avvio subprocess trascrizione...")
                self.progress.emit(f"Trascrizione in corso (modello: {model_name})...")
                result = subprocess.run(
                    [python_exe, "-c", script],
                    capture_output=True,
                    text=True,
                    timeout=600,
                    env={**os.environ, "HF_HUB_DISABLE_XET": "1"},
                )
                _dbg(f"subprocess returncode: {result.returncode}")
                if result.returncode == 0:
                    text = result.stdout.strip()
                    _dbg(f"trascrizione OK: {len(text.split())} parole")
                    self.progress.emit(f"Trascrizione completata — {len(text.split())} parole")
                    self.finished.emit(text)
                else:
                    err = result.stderr.strip()
                    _dbg(f"subprocess ERRORE: {err}")
                    self.error.emit(f"Errore trascrizione:\n{err}")

        except Exception as e:
            import traceback
            err = traceback.format_exc()
            _dbg(f"ERRORE: {err}")
            self.error.emit(f"Errore trascrizione:\n{str(e)}")


class ReportWorker(QThread):
    progress = pyqtSignal(str)
    finished = pyqtSignal(str)
    error    = pyqtSignal(str)

    SYSTEM_PROMPT = (
        "Sei un assistente specializzato nell'analisi di trascrizioni di riunioni aziendali. "
        "Produci report precisi e professionali in italiano. "
        "Rispondi SOLO con il report, senza premesse o commenti aggiuntivi."
    )

    SCHEMAS = {
        "📋  Strutturato": """Analizza la seguente trascrizione e genera un report strutturato:

## 📋 SOMMARIO ESECUTIVO
(3-5 frasi sui punti principali)

## 🎯 DECISIONI PRESE
(elenco numerato)

## ✅ ACTION ITEMS
(tabella: Azione | Responsabile | Scadenza)

## 💡 PUNTI CHIAVE DISCUSSI

## ⚠️ RISCHI / PROBLEMI EMERSI

## 📌 NOTE AGGIUNTIVE

---
TRASCRIZIONE:
{transcript}""",

        "🆓  Libero": """Analizza la seguente trascrizione di una riunione aziendale e genera un report completo nel formato che ritieni più appropriato per il contenuto. Scegli liberamente la struttura, le sezioni e il livello di dettaglio più adatti.

TRASCRIZIONE:
{transcript}""",

        "📝  Verbale formale": """Redigi un verbale formale di riunione basandoti sulla seguente trascrizione. Usa uno stile burocratico e formale. Includi:

**VERBALE DI RIUNIONE**

**Data e ora:** (estrai dalla trascrizione o indica "come da convocazione")
**Partecipanti:** (elenca i nomi citati)
**Ordine del giorno:**
**Svolgimento della riunione:**
**Delibere:**
**Incarichi assegnati:**
**Data prossima riunione:** (se menzionata)

Il Segretario

---
TRASCRIZIONE:
{transcript}""",

        "⚡  Solo Action Items": """Dalla seguente trascrizione estrai ESCLUSIVAMENTE gli action items, le attività assegnate e le scadenze menzionate. Formato tabella:

| # | Azione | Responsabile | Scadenza | Priorità |
|---|--------|-------------|----------|----------|

Se non è indicato un responsabile o una scadenza, scrivi "Da definire".
Aggiungi una riga di **TOTALE** con il conteggio degli action items.

---
TRASCRIZIONE:
{transcript}""",

        "📊  Executive Summary": """Genera un executive summary conciso (massimo 10 righe) della seguente riunione, adatto ad essere letto da un dirigente in 30 secondi. Includi solo: obiettivo della riunione, decisioni chiave, prossimi passi critici.

TRASCRIZIONE:
{transcript}""",
    }

    def __init__(self, transcript, cfg, schema=None):
        super().__init__()
        self.transcript = transcript
        self.cfg        = cfg
        self.schema     = schema or "📋  Strutturato"

    def run(self):
        import requests as req
        try:
            self.progress.emit("Connessione a LM Studio...")
            template = self.SCHEMAS.get(self.schema, self.SCHEMAS["📋  Strutturato"])
            user_content = self.SYSTEM_PROMPT + "\n\n" + template.format(transcript=self.transcript)
            payload = {
                "model": self.cfg["lm_model"],
                "messages": [
                    {"role": "user", "content": user_content},
                ],
                "max_tokens": 2000,
                "temperature": 0.3,
                "stream": False,
            }
            self.progress.emit("Generazione report in corso...")
            resp = req.post(
                self.cfg["lm_url"],
                headers={"Content-Type": "application/json", "Accept": "application/json"},
                data=json.dumps(payload),
                timeout=600,
            )
            _dbg(f"LM Studio status: {resp.status_code}")
            resp.raise_for_status()
            data = resp.json()
            if "choices" in data:
                report = data["choices"][0]["message"]["content"].strip()
            else:
                report = str(data)
            self.progress.emit("Report generato con successo!")
            self.finished.emit(report)
        except req.exceptions.ConnectionError:
            self.error.emit(
                "Impossibile connettersi a LM Studio.\n"
                "Verifica che il server locale sia attivo su porta 1234."
            )
        except Exception as e:
            self.error.emit(str(e))


# ──────────────────────────────────────────────────────
# WIDGET LIVELLO AUDIO
# ──────────────────────────────────────────────────────
class AudioLevelBar(QWidget):
    def __init__(self):
        super().__init__()
        self.level = 0.0
        self.setFixedHeight(8)
        self.setMinimumWidth(200)

    def set_level(self, v):
        self.level = v
        self.update()

    def paintEvent(self, e):
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        w, h = self.width(), self.height()
        # background
        p.setBrush(QColor("#1E293B"))
        p.setPen(Qt.PenStyle.NoPen)
        p.drawRoundedRect(0, 0, w, h, 4, 4)
        # fill
        filled = int(w * self.level)
        if filled > 0:
            grad = QLinearGradient(0, 0, w, 0)
            grad.setColorAt(0.0, QColor("#2563EB"))
            grad.setColorAt(0.6, QColor("#7C3AED"))
            grad.setColorAt(1.0, QColor("#EF4444"))
            p.setBrush(QBrush(grad))
            p.drawRoundedRect(0, 0, filled, h, 4, 4)
        p.end()


# ──────────────────────────────────────────────────────
# SCHEDA REGISTRAZIONE
# ──────────────────────────────────────────────────────
class RecordTab(QWidget):
    session_done = pyqtSignal(dict)   # emette dati sessione completata

    def __init__(self, cfg_getter):
        super().__init__()
        self.cfg_getter = cfg_getter
        self.record_worker = None
        self.transcribe_worker = None
        self.report_worker = None
        self.wav_path = None
        self.transcript = ""
        self.elapsed = 0
        self.timer = QTimer()
        self.timer.timeout.connect(self._tick)
        self._build()

    def _build(self):
        root = QVBoxLayout(self)
        root.setSpacing(16)
        root.setContentsMargins(24, 20, 24, 20)

        # ── Header ──
        hdr = QHBoxLayout()
        lbl = QLabel("Registrazione riunione")
        lbl.setObjectName("title")
        hdr.addWidget(lbl)
        hdr.addStretch()
        self.lbl_time = QLabel("00:00")
        self.lbl_time.setFont(QFont("Courier New", 20, QFont.Weight.Bold))
        self.lbl_time.setStyleSheet("color: #60A5FA;")
        hdr.addWidget(self.lbl_time)
        root.addLayout(hdr)

        # ── Livello audio ──
        self.audio_bar = AudioLevelBar()
        root.addWidget(self.audio_bar)

        # ── Pulsanti principali ──
        btn_row = QHBoxLayout()
        btn_row.setSpacing(12)

        self.btn_record = QPushButton("⏺  Registra")
        self.btn_record.setFixedHeight(52)
        self.btn_record.setObjectName("btnSuccess")
        self.btn_record.clicked.connect(self._toggle_record)
        btn_row.addWidget(self.btn_record)

        self.btn_load = QPushButton("📂  Carica file audio")
        self.btn_load.setFixedHeight(52)
        self.btn_load.setObjectName("btnSecondary")
        self.btn_load.clicked.connect(self._load_file)
        btn_row.addWidget(self.btn_load)

        root.addLayout(btn_row)

        # ── Status ──
        self.lbl_status = QLabel("Pronto")
        self.lbl_status.setObjectName("subtitle")
        root.addWidget(self.lbl_status)

        # ── Progress ──
        self.progress = QProgressBar()
        self.progress.setRange(0, 0)
        self.progress.setVisible(False)
        root.addWidget(self.progress)

        # ── Splitter trascrizione / report ──
        splitter = QSplitter(Qt.Orientation.Vertical)

        # Trascrizione
        grp_tr = QGroupBox("📝  Trascrizione")
        vtr = QVBoxLayout(grp_tr)
        self.txt_transcript = QTextEdit()
        self.txt_transcript.setPlaceholderText("La trascrizione apparirà qui dopo la registrazione o incolla il testo...")
        self.txt_transcript.setMinimumHeight(120)
        self.txt_transcript.textChanged.connect(self._on_transcript_text_changed)
        vtr.addWidget(self.txt_transcript)
        btn_tr_row = QHBoxLayout()
        self.btn_transcribe = QPushButton("🎙  Trascrivi ora")
        self.btn_transcribe.setObjectName("btnSecondary")
        self.btn_transcribe.setEnabled(False)
        self.btn_transcribe.clicked.connect(self._do_transcribe)
        btn_tr_row.addWidget(self.btn_transcribe)
        btn_tr_row.addStretch()
        vtr.addLayout(btn_tr_row)
        splitter.addWidget(grp_tr)

        # Report
        grp_rep = QGroupBox("📊  Report")
        vrep = QVBoxLayout(grp_rep)
        self.txt_report = QTextEdit()
        self.txt_report.setPlaceholderText("Il report generato da LM Studio apparirà qui...")
        self.txt_report.setMinimumHeight(120)
        vrep.addWidget(self.txt_report)
        btn_rep_row = QHBoxLayout()
        self.cmb_schema = QComboBox()
        self.cmb_schema.addItems(list(ReportWorker.SCHEMAS.keys()))
        self.cmb_schema.setFixedHeight(36)
        self.cmb_schema.setToolTip("Seleziona il formato del report")
        btn_rep_row.addWidget(self.cmb_schema)
        self.btn_report = QPushButton("🤖  Genera report")
        self.btn_report.setEnabled(False)
        self.btn_report.clicked.connect(self._do_report)
        btn_rep_row.addWidget(self.btn_report)
        self.btn_save = QPushButton("💾  Salva")
        self.btn_save.setObjectName("btnSecondary")
        self.btn_save.setEnabled(False)
        self.btn_save.clicked.connect(self._save_all)
        btn_rep_row.addWidget(self.btn_save)
        btn_rep_row.addStretch()
        vrep.addLayout(btn_rep_row)
        splitter.addWidget(grp_rep)

        root.addWidget(splitter)

    # ── Recording ──────────────────────────────────────
    def _toggle_record(self):
        if self.record_worker and self.record_worker.isRunning():
            self._stop_recording()
        else:
            self._start_recording()

    def _start_recording(self):
        cfg = self.cfg_getter()
        self.record_worker = RecordWorker(cfg["output_dir"], cfg)
        self.record_worker.finished.connect(self._on_recording_done)
        self.record_worker.error.connect(self._on_error)
        self.record_worker.level.connect(self.audio_bar.set_level)
        self.record_worker.start()
        self.elapsed = 0
        self.timer.start(1000)
        self.btn_record.setText("⏹  Stop")
        self.btn_record.setObjectName("btnDanger")
        self.btn_record.setStyle(self.btn_record.style())
        self.btn_load.setEnabled(False)
        self._set_status("🔴  Registrazione in corso...", "recording")
        self.progress.setVisible(False)

    def _stop_recording(self):
        if self.record_worker:
            self.record_worker.stop()
        self.timer.stop()
        self.btn_record.setEnabled(False)
        self._set_status("Salvataggio audio...", "subtitle")

    def _on_recording_done(self, path):
        _dbg("_on_recording_done: " + path)
        self.wav_path = path
        self.audio_bar.set_level(0)
        self.btn_record.setText("⏺  Registra")
        self.btn_record.setObjectName("btnSuccess")
        self.btn_record.setStyle(self.btn_record.style())
        self.btn_record.setEnabled(True)
        self.btn_load.setEnabled(True)
        self.btn_transcribe.setEnabled(True)
        dur = self.elapsed
        self._set_status(f"✅  Audio salvato — {dur}s  |  {Path(path).name}", "success")
        _dbg("avvio trascrizione automatica")
        self._do_transcribe()

    def _tick(self):
        self.elapsed += 1
        m, s = divmod(self.elapsed, 60)
        self.lbl_time.setText(f"{m:02d}:{s:02d}")

    # ── Load file ──────────────────────────────────────
    def _load_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Carica file audio", str(Path.home()),
            "Audio (*.wav *.mp3 *.m4a *.ogg *.flac);;Tutti (*.*)"
        )
        if path:
            self.wav_path = path
            self.btn_transcribe.setEnabled(True)
            self._set_status(f"📂  File caricato: {Path(path).name}", "success")
            self._do_transcribe()

    # ── Transcribe ─────────────────────────────────────
    def _do_transcribe(self):
        if not self.wav_path:
            return
        cfg = self.cfg_getter()
        self.transcribe_worker = TranscribeWorker(self.wav_path, cfg)
        self.transcribe_worker.progress.connect(self._set_status_plain)
        self.transcribe_worker.finished.connect(self._on_transcript_done)
        self.transcribe_worker.error.connect(self._on_error)
        self.transcribe_worker.start()
        self.progress.setVisible(True)
        self.btn_transcribe.setEnabled(False)
        self.btn_report.setEnabled(False)

    def _on_transcript_done(self, text):
        self.transcript = text
        self.txt_transcript.setPlainText(text)
        self.progress.setVisible(False)
        self.btn_transcribe.setEnabled(True)
        self.btn_report.setEnabled(True)
        self._set_status("✅  Trascrizione completata", "success")

    # ── Report ─────────────────────────────────────────
    def _do_report(self):
        text = self.txt_transcript.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Attenzione", "Nessuna trascrizione disponibile.")
            return
        cfg = self.cfg_getter()
        self.report_worker = ReportWorker(text, cfg, self.cmb_schema.currentText())
        self.report_worker.progress.connect(self._set_status_plain)
        self.report_worker.finished.connect(self._on_report_done)
        self.report_worker.error.connect(self._on_error)
        self.report_worker.start()
        self.progress.setVisible(True)
        self.btn_report.setEnabled(False)

    def _on_report_done(self, report):
        self.txt_report.setPlainText(report)
        self.progress.setVisible(False)
        self.btn_report.setEnabled(True)
        self.btn_save.setEnabled(True)
        self._set_status("✅  Report generato con successo", "success")

    # ── Save ───────────────────────────────────────────
    def _save_all(self):
        cfg = self.cfg_getter()
        out = Path(cfg["output_dir"])
        out.mkdir(parents=True, exist_ok=True)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        base = f"riunione_{ts}"
        if self.wav_path:
            base = Path(self.wav_path).stem

        tr_path  = out / f"{base}_trascrizione.txt"
        rep_path = out / f"{base}_report.md"

        tr_path.write_text(self.txt_transcript.toPlainText(), encoding="utf-8")
        header = (
            f"# Report Riunione\n"
            f"**Data:** {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n---\n\n"
        )
        rep_path.write_text(header + self.txt_report.toPlainText(), encoding="utf-8")

        self.session_done.emit({
            "date":       datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),
            "audio":      str(self.wav_path) if self.wav_path else "",
            "transcript": str(tr_path),
            "report":     str(rep_path),
        })
        self._set_status(f"💾  Salvato in {out}", "success")
        QMessageBox.information(self, "Salvato", f"File salvati in:\n{out}")

    # ── Auto-enable report on paste/edit ─────────────────
    def _on_transcript_text_changed(self):
        """Abilita il bottone Genera report quando il campo trascrizione contiene testo."""
        has_text = bool(self.txt_transcript.toPlainText().strip())
        self.btn_report.setEnabled(has_text)

    # ── Helpers ────────────────────────────────────────
    def _set_status(self, msg, obj="subtitle"):
        self.lbl_status.setText(msg)
        self.lbl_status.setObjectName(obj)
        self.lbl_status.setStyle(self.lbl_status.style())

    def _set_status_plain(self, msg):
        self._set_status(msg, "subtitle")

    def _on_error(self, msg):
        self.progress.setVisible(False)
        self._set_status(f"❌  {msg}", "warning")
        QMessageBox.critical(self, "Errore", msg)


# ──────────────────────────────────────────────────────
# SCHEDA REPORT SALVATI
# ──────────────────────────────────────────────────────
class ReportsTab(QWidget):
    def __init__(self, cfg_getter):
        super().__init__()
        self.cfg_getter = cfg_getter
        self.current_path = None
        self._build()

    def _build(self):
        root = QHBoxLayout(self)
        root.setSpacing(12)
        root.setContentsMargins(24, 20, 24, 20)

        # Lista
        left = QVBoxLayout()
        lbl = QLabel("Report salvati")
        lbl.setObjectName("title")
        left.addWidget(lbl)

        self.list_reports = QListWidget()
        self.list_reports.currentItemChanged.connect(self._on_select)
        left.addWidget(self.list_reports)

        btn_row = QHBoxLayout()
        btn_refresh = QPushButton("🔄  Aggiorna")
        btn_refresh.setObjectName("btnSecondary")
        btn_refresh.clicked.connect(self.refresh)
        btn_row.addWidget(btn_refresh)

        btn_open = QPushButton("📂  Apri cartella")
        btn_open.setObjectName("btnSecondary")
        btn_open.clicked.connect(self._open_folder)
        btn_row.addWidget(btn_open)
        left.addLayout(btn_row)

        left_w = QWidget()
        left_w.setLayout(left)
        left_w.setFixedWidth(300)
        root.addWidget(left_w)

        # Visualizzatore
        right = QVBoxLayout()
        right.setSpacing(8)

        # Header con nome file e pulsanti
        top_row = QHBoxLayout()
        self.lbl_filename = QLabel("Seleziona un report")
        self.lbl_filename.setObjectName("subtitle")
        top_row.addWidget(self.lbl_filename)
        top_row.addStretch()

        btn_word = QPushButton("📄  Esporta Word")
        btn_word.setObjectName("btnSecondary")
        btn_word.setFixedHeight(36)
        btn_word.clicked.connect(self._export_word)
        top_row.addWidget(btn_word)

        btn_print = QPushButton("🖨  Stampa")
        btn_print.setObjectName("btnSecondary")
        btn_print.setFixedHeight(36)
        btn_print.clicked.connect(self._print_report)
        top_row.addWidget(btn_print)

        right.addLayout(top_row)

        self.txt_view = QTextEdit()
        self.txt_view.setReadOnly(True)
        right.addWidget(self.txt_view)
        right_w = QWidget()
        right_w.setLayout(right)
        root.addWidget(right_w)

        self.refresh()

    def refresh(self):
        self.list_reports.clear()
        cfg = self.cfg_getter()
        out = Path(cfg["output_dir"])
        if not out.exists():
            return
        reports = sorted(out.glob("*_report.md"), reverse=True)
        for r in reports:
            item = QListWidgetItem(f"📄  {r.stem.replace('_report','')}")
            item.setData(Qt.ItemDataRole.UserRole, str(r))
            self.list_reports.addItem(item)

    def _on_select(self, item):
        if not item:
            return
        path = item.data(Qt.ItemDataRole.UserRole)
        self.current_path = path
        self.lbl_filename.setText(Path(path).name)
        try:
            self.txt_view.setPlainText(Path(path).read_text(encoding="utf-8"))
        except Exception as e:
            self.txt_view.setPlainText(f"Errore lettura: {e}")

    def _export_word(self):
        if not hasattr(self, 'current_path') or not self.current_path:
            QMessageBox.warning(self, "Attenzione", "Seleziona prima un report dalla lista.")
            return
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # Installa python-docx se non presente
            import subprocess, sys
            subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
            try:
                from docx import Document as DocxDocument
                from docx.shared import Pt, RGBColor
            except ImportError:
                QMessageBox.critical(self, "Errore",
                    "Impossibile installare python-docx.\n"
                    "Esegui manualmente: pip install python-docx")
                return

        try:
            text = Path(self.current_path).read_text(encoding="utf-8")
            doc = DocxDocument()

            # Stile base
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            for line in text.split('\n'):
                line = line.rstrip()
                if line.startswith('## '):
                    p = doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    p = doc.add_heading(line[2:], level=1)
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('|') and '|' in line[1:]:
                    # Riga tabella Markdown — aggiunge come testo normale
                    if not all(c in '|-: ' for c in line):
                        doc.add_paragraph(line)
                elif line == '---':
                    doc.add_paragraph('─' * 50)
                elif line == '':
                    doc.add_paragraph('')
                else:
                    doc.add_paragraph(line)

            # Proponi nome file
            default_name = Path(self.current_path).stem.replace('_report', '') + '_report.docx'
            save_path, _ = QFileDialog.getSaveFileName(
                self, "Salva documento Word",
                str(Path(self.current_path).parent / default_name),
                "Word Document (*.docx)"
            )
            if save_path:
                doc.save(save_path)
                QMessageBox.information(self, "Esportato", f"Report salvato in:\n{save_path}")

                # Apri il file automaticamente
                if sys.platform == "darwin":
                    subprocess.run(["open", save_path])
                elif sys.platform == "win32":
                    subprocess.run(["start", "", save_path], shell=True)

        except Exception as e:
            QMessageBox.critical(self, "Errore esportazione", str(e))

    def _print_report(self):
        if not hasattr(self, 'current_path') or not self.current_path:
            QMessageBox.warning(self, "Attenzione", "Seleziona prima un report dalla lista.")
            return
        try:
            from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
            from PyQt6.QtGui import QTextDocument

            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            dialog = QPrintDialog(printer, self)
            if dialog.exec() == QPrintDialog.DialogCode.Accepted:
                text = self.txt_view.toPlainText()
                doc = QTextDocument()
                # Formattazione base per la stampa
                html = "<html><body style='font-family:Arial;font-size:11pt;'>"
                for line in text.split('\n'):
                    line = line.rstrip()
                    if line.startswith('## '):
                        html += f"<h2>{line[3:]}</h2>"
                    elif line.startswith('# '):
                        html += f"<h1>{line[2:]}</h1>"
                    elif line.startswith('- ') or line.startswith('* '):
                        html += f"<li>{line[2:]}</li>"
                    elif line == '---':
                        html += "<hr>"
                    elif line == '':
                        html += "<br>"
                    else:
                        html += f"<p>{line}</p>"
                html += "</body></html>"
                doc.setHtml(html)
                doc.print(printer)
        except Exception as e:
            # Fallback: apri il file con l'app di sistema che gestisce la stampa
            try:
                if sys.platform == "win32":
                    subprocess.run(["notepad", "/p", self.current_path], shell=True)
                elif sys.platform == "darwin":
                    subprocess.run(["lpr", self.current_path])
                else:
                    subprocess.run(["lpr", self.current_path])
            except Exception as e2:
                QMessageBox.critical(self, "Errore stampa", str(e2))

    def _open_folder(self):
        cfg = self.cfg_getter()
        path = cfg["output_dir"]
        if sys.platform == "darwin":
            subprocess.run(["open", path])
        elif sys.platform == "win32":
            subprocess.run(["explorer", path])
        else:
            subprocess.run(["xdg-open", path])

    def add_session(self, data):
        self.refresh()


# ──────────────────────────────────────────────────────
# SCHEDA IMPOSTAZIONI
# ──────────────────────────────────────────────────────
class SettingsTab(QWidget):
    config_changed = pyqtSignal(dict)

    def __init__(self):
        super().__init__()
        self.cfg = load_config()
        self._build()

    def _build(self):
        root = QVBoxLayout(self)
        root.setSpacing(20)
        root.setContentsMargins(32, 24, 32, 24)

        lbl = QLabel("Impostazioni")
        lbl.setObjectName("title")
        root.addWidget(lbl)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        inner = QWidget()
        vlay = QVBoxLayout(inner)
        vlay.setSpacing(20)

        # ── LM Studio ──
        grp_lm = QGroupBox("LM Studio")
        flm = QFormLayout(grp_lm)
        flm.setSpacing(12)
        self.inp_url = QLineEdit(self.cfg["lm_url"])
        flm.addRow("URL Server:", self.inp_url)
        self.inp_model = QLineEdit(self.cfg["lm_model"])
        flm.addRow("Nome modello:", self.inp_model)
        self.btn_test = QPushButton("🔌  Testa connessione")
        self.btn_test.setObjectName("btnSecondary")
        self.btn_test.clicked.connect(self._test_connection)
        self.lbl_conn = QLabel("")
        flm.addRow(self.btn_test, self.lbl_conn)
        vlay.addWidget(grp_lm)

        # ── Whisper ──
        grp_w = QGroupBox("Whisper — Trascrizione")
        fw = QFormLayout(grp_w)
        fw.setSpacing(12)
        self.cmb_model = QComboBox()
        self.cmb_model.addItems(["tiny", "base", "small", "medium", "large-v3"])
        self.cmb_model.setCurrentText(self.cfg["whisper_model"])
        fw.addRow("Modello:", self.cmb_model)
        self.cmb_lang = QComboBox()
        self.cmb_lang.addItems(["it", "en", "fr", "de", "es", "auto"])
        self.cmb_lang.setCurrentText(self.cfg["language"])
        fw.addRow("Lingua:", self.cmb_lang)
        self.chk_diarize = QCheckBox("Identifica speaker (diarizzazione)")
        self.chk_diarize.setChecked(self.cfg.get("diarization", True))
        self.chk_diarize.setToolTip(
            "Se attivo, la trascrizione identifica i diversi speaker.\n"
            "Con pyannote.audio + HF_TOKEN: diarizzazione avanzata per timbro vocale.\n"
            "Senza pyannote: euristica basata sulle pause tra i segmenti."
        )
        fw.addRow("", self.chk_diarize)
        vlay.addWidget(grp_w)

        # ── Output ──
        grp_out = QGroupBox("Output")
        fo = QFormLayout(grp_out)
        fo.setSpacing(12)
        row_out = QHBoxLayout()
        self.inp_output = QLineEdit(self.cfg["output_dir"])
        row_out.addWidget(self.inp_output)
        btn_browse = QPushButton("📁")
        btn_browse.setFixedWidth(40)
        btn_browse.setObjectName("btnSecondary")
        btn_browse.clicked.connect(self._browse_output)
        row_out.addWidget(btn_browse)
        fo.addRow("Cartella output:", row_out)
        vlay.addWidget(grp_out)

        vlay.addStretch()
        scroll.setWidget(inner)
        root.addWidget(scroll)

        # Save button
        btn_save = QPushButton("💾  Salva impostazioni")
        btn_save.setFixedHeight(44)
        btn_save.clicked.connect(self._save)
        root.addWidget(btn_save)

    def _test_connection(self):
        import requests as req
        url = self.inp_url.text().replace("/chat/completions", "/models")
        try:
            r = req.get(url, timeout=5)
            if r.status_code == 200:
                self.lbl_conn.setText("✅  Connesso")
                self.lbl_conn.setObjectName("success")
            else:
                self.lbl_conn.setText(f"⚠️  HTTP {r.status_code}")
                self.lbl_conn.setObjectName("warning")
        except Exception:
            self.lbl_conn.setText("❌  Non raggiungibile")
            self.lbl_conn.setObjectName("warning")
        self.lbl_conn.setStyle(self.lbl_conn.style())

    def _browse_output(self):
        path = QFileDialog.getExistingDirectory(self, "Scegli cartella output", self.inp_output.text())
        if path:
            self.inp_output.setText(path)

    def _save(self):
        self.cfg.update({
            "lm_url":        self.inp_url.text().strip(),
            "lm_model":      self.inp_model.text().strip(),
            "whisper_model": self.cmb_model.currentText(),
            "language":      self.cmb_lang.currentText(),
            "output_dir":    self.inp_output.text().strip(),
            "diarization":   self.chk_diarize.isChecked(),
        })
        save_config(self.cfg)
        self.config_changed.emit(self.cfg)
        QMessageBox.information(self, "Salvato", "Impostazioni salvate correttamente.")

    def get_config(self):
        return self.cfg.copy()


# ──────────────────────────────────────────────────────
# FINESTRA CHATBOT
# ──────────────────────────────────────────────────────
class ChatWorker(QThread):
    token    = pyqtSignal(str)
    finished = pyqtSignal()
    error    = pyqtSignal(str)

    TONES = {
        "🏢  Aziendale formale": (
            "Sei un assistente aziendale professionale. "
            "Rispondi SEMPRE in italiano, con tono formale, preciso e autorevole. "
            "Usa un linguaggio professionale, evita colloquialismi. "
            "Sii conciso ed esaustivo."
        ),
        "💬  Informale": (
            "Sei un assistente simpatico e disponibile. "
            "Rispondi SEMPRE in italiano, con tono amichevole e diretto. "
            "Puoi usare un linguaggio colloquiale, ma rimani sempre corretto e utile. "
            "Sii chiaro e vai al punto."
        ),
        "🎓  Tecnico": (
            "Sei un esperto tecnico altamente specializzato. "
            "Rispondi SEMPRE in italiano, con linguaggio tecnico preciso e dettagliato. "
            "Fornisci spiegazioni approfondite, usa terminologia specialistica quando appropriato. "
            "Struttura le risposte in modo logico e metodico."
        ),
        "📝  Sintetico": (
            "Sei un assistente che valorizza la brevità. "
            "Rispondi SEMPRE in italiano, nel modo più conciso possibile. "
            "Usa elenchi puntati quando utile. "
            "Niente introduzioni o conclusioni inutili — solo l'essenziale."
        ),
        "🧠  Analitico": (
            "Sei un consulente analitico e strategico. "
            "Rispondi SEMPRE in italiano, analizzando ogni problema da più angolazioni. "
            "Evidenzia pro e contro, fornisci ragionamenti strutturati e conclusioni chiare. "
            "Basa le risposte su logica e dati quando disponibili."
        ),
    }

    def __init__(self, messages, cfg, tone_key=None):
        super().__init__()
        self.messages  = messages
        self.cfg       = cfg
        self.tone_key  = tone_key or "🏢  Aziendale formale"

    def run(self):
        import requests as req
        try:
            system = self.TONES.get(self.tone_key, list(self.TONES.values())[0])

            api_messages = []
            for i, msg in enumerate(self.messages):
                if i == 0 and msg["role"] == "user":
                    api_messages.append({
                        "role": "user",
                        "content": system + "\n\n" + msg["content"]
                    })
                else:
                    api_messages.append(msg)

            payload = {
                "model": self.cfg["lm_model"],
                "messages": api_messages,
                "max_tokens": 1500,
                "temperature": 0.4,
                "stream": False,
            }
            resp = req.post(
                self.cfg["lm_url"],
                headers={"Content-Type": "application/json"},
                data=json.dumps(payload),
                timeout=300,
            )
            resp.raise_for_status()
            data = resp.json()
            text = data["choices"][0]["message"]["content"].strip() if "choices" in data else str(data)
            self.token.emit(text)
            self.finished.emit()
        except Exception as e:
            self.error.emit(str(e))


class ChatWindow(QDialog):
    WELCOME = (
        "Buongiorno. Sono il suo assistente aziendale.\n"
        "Come posso aiutarla oggi? Posso analizzare trascrizioni, "
        "redigere documenti aziendali, o rispondere a qualsiasi domanda professionale."
    )

    def __init__(self, cfg_getter, get_transcript=None, parent=None):
        super().__init__(parent)
        self.cfg_getter     = cfg_getter
        self.get_transcript = get_transcript
        self.history        = []          # conversazione corrente
        self.all_chats      = []          # lista di chat salvate [{title, history}]
        self.worker         = None
        self.sidebar_visible = True
        self.setWindowTitle("Assistente Aziendale")
        self.setMinimumSize(900, 650)
        self.resize(1100, 720)
        self.setStyleSheet("""
            QDialog { background-color: #212121; }
            QWidget#sidebar {
                background-color: #171717;
                border-right: 1px solid #2A2A2A;
            }
            QWidget#chat_area { background-color: #212121; }
            QScrollArea { background-color: #212121; border: none; }
            QWidget#msg_container { background-color: #212121; }
            QListWidget {
                background-color: #171717;
                border: none;
                color: #ECECEC;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 13px;
                outline: none;
            }
            QListWidget::item {
                padding: 10px 14px;
                border-radius: 6px;
                margin: 2px 6px;
            }
            QListWidget::item:selected {
                background-color: #2A2A2A;
                color: #FFFFFF;
            }
            QListWidget::item:hover:!selected {
                background-color: #222222;
            }
            QTextEdit#msg_input {
                background-color: #2F2F2F;
                border: 1px solid #3F3F3F;
                border-radius: 12px;
                color: #ECECEC;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 14px;
                padding: 10px 14px;
            }
            QTextEdit#msg_input:focus { border: 1px solid #555555; }
            QPushButton#send_btn {
                background-color: #ECECEC;
                color: #212121;
                border: none;
                border-radius: 8px;
                font-size: 18px;
                font-weight: bold;
            }
            QPushButton#send_btn:hover { background-color: #FFFFFF; }
            QPushButton#send_btn:disabled { background-color: #3F3F3F; color: #555555; }
            QPushButton#sidebar_btn {
                background-color: transparent;
                border: none;
                color: #888888;
                font-size: 18px;
                padding: 4px;
            }
            QPushButton#sidebar_btn:hover { color: #ECECEC; }
            QPushButton#new_chat_btn {
                background-color: #2A2A2A;
                border: 1px solid #3A3A3A;
                border-radius: 8px;
                color: #ECECEC;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 13px;
                padding: 8px 12px;
            }
            QPushButton#new_chat_btn:hover { background-color: #333333; }
            QPushButton#icon_btn {
                background-color: transparent;
                border: none;
                color: #888888;
                font-size: 12px;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
            QPushButton#icon_btn:hover { color: #ECECEC; }
            QComboBox {
                background-color: #2F2F2F;
                border: 1px solid #3F3F3F;
                border-radius: 8px;
                color: #ECECEC;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 12px;
                padding: 4px 10px;
                min-width: 170px;
            }
            QComboBox:hover { border: 1px solid #555555; }
            QComboBox::drop-down { border: none; width: 18px; }
            QComboBox QAbstractItemView {
                background-color: #2F2F2F;
                border: 1px solid #3F3F3F;
                color: #ECECEC;
                selection-background-color: #19C37D;
                selection-color: #000000;
                font-size: 13px;
            }
            QScrollBar:vertical {
                background: #212121; width: 6px; border-radius: 3px;
            }
            QScrollBar::handle:vertical {
                background: #3F3F3F; border-radius: 3px; min-height: 30px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; }
        """)
        self._build()

    def _build(self):
        root = QHBoxLayout(self)
        root.setSpacing(0)
        root.setContentsMargins(0, 0, 0, 0)

        # ══ SIDEBAR ══════════════════════════════════════
        self.sidebar = QWidget()
        self.sidebar.setObjectName("sidebar")
        self.sidebar.setFixedWidth(240)
        sl = QVBoxLayout(self.sidebar)
        sl.setContentsMargins(0, 12, 0, 12)
        sl.setSpacing(6)

        # Header sidebar
        sh = QHBoxLayout()
        sh.setContentsMargins(12, 0, 12, 0)
        lbl_s = QLabel("Meeting Recorder")
        lbl_s.setStyleSheet("color: #ECECEC; font-size: 13px; font-weight: 600; font-family: 'Segoe UI', Arial;")
        sh.addWidget(lbl_s)
        sh.addStretch()
        btn_toggle = QPushButton("◀")
        btn_toggle.setObjectName("sidebar_btn")
        btn_toggle.setFixedSize(28, 28)
        btn_toggle.setToolTip("Nascondi sidebar")
        btn_toggle.clicked.connect(self._toggle_sidebar)
        sh.addWidget(btn_toggle)
        self.btn_toggle = btn_toggle
        sl.addLayout(sh)

        # Nuova chat
        btn_new = QPushButton("✏  Nuova chat")
        btn_new.setObjectName("new_chat_btn")
        btn_new.setFixedHeight(38)
        btn_new.clicked.connect(self._new_chat)
        btn_new.setContentsMargins(12, 0, 12, 0)
        sl.addWidget(btn_new)

        # Allega trascrizione
        btn_attach = QPushButton("📋  Allega trascrizione")
        btn_attach.setObjectName("new_chat_btn")
        btn_attach.setFixedHeight(34)
        btn_attach.clicked.connect(self._attach_transcript)
        sl.addWidget(btn_attach)

        # Pulisci chat
        btn_clear_side = QPushButton("🗑  Pulisci chat")
        btn_clear_side.setObjectName("new_chat_btn")
        btn_clear_side.setFixedHeight(34)
        btn_clear_side.clicked.connect(self._clear_chat)
        sl.addWidget(btn_clear_side)

        self.lbl_attach = QLabel("")
        self.lbl_attach.setStyleSheet("color: #10B981; font-size: 11px; font-family: 'Segoe UI', Arial; padding: 0 14px;")
        self.lbl_attach.setWordWrap(True)
        sl.addWidget(self.lbl_attach)

        sl.addSpacing(8)

        # Lista cronologia
        lbl_hist = QLabel("Cronologia")
        lbl_hist.setStyleSheet("color: #555555; font-size: 11px; font-family: 'Segoe UI'; padding: 0 14px;")
        sl.addWidget(lbl_hist)

        self.list_history = QListWidget()
        self.list_history.setSpacing(0)
        self.list_history.itemClicked.connect(self._load_chat)
        self.list_history.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.list_history.customContextMenuRequested.connect(self._show_context_menu)
        sl.addWidget(self.list_history)

        # Pulsante elimina chat selezionata
        btn_del_sel = QPushButton("🗑  Elimina chat selezionata")
        btn_del_sel.setObjectName("new_chat_btn")
        btn_del_sel.setFixedHeight(32)
        btn_del_sel.setStyleSheet("""
            QPushButton {
                background-color: #2A2A2A;
                border: 1px solid #3A3A3A;
                border-radius: 8px;
                color: #EF4444;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 12px;
                padding: 4px 10px;
            }
            QPushButton:hover { background-color: #3A1A1A; }
        """)
        btn_del_sel.clicked.connect(self._delete_selected_chat)
        sl.addWidget(btn_del_sel)

        sl.addStretch()

        # Credits in fondo alla sidebar
        lbl_credit = QLabel("Sviluppato da\nMarco Bonometti")
        lbl_credit.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl_credit.setStyleSheet("""
            color: #333333;
            font-size: 11px;
            font-family: 'Segoe UI', Arial, sans-serif;
            padding: 8px 0;
        """)
        sl.addWidget(lbl_credit)

        root.addWidget(self.sidebar)

        # ══ AREA CHAT ════════════════════════════════════
        chat_widget = QWidget()
        chat_widget.setObjectName("chat_area")
        cl = QVBoxLayout(chat_widget)
        cl.setSpacing(0)
        cl.setContentsMargins(0, 0, 0, 0)

        # Header chat
        header = QFrame()
        header.setFixedHeight(52)
        header.setStyleSheet("background-color: #212121; border-bottom: 1px solid #2A2A2A;")
        hl = QHBoxLayout(header)
        hl.setContentsMargins(16, 0, 16, 0)
        hl.setSpacing(12)

        # Bottone mostra/nascondi sidebar (quando è nascosta)
        self.btn_show_sidebar = QPushButton("▶")
        self.btn_show_sidebar.setObjectName("sidebar_btn")
        self.btn_show_sidebar.setFixedSize(28, 28)
        self.btn_show_sidebar.setToolTip("Mostra sidebar")
        self.btn_show_sidebar.clicked.connect(self._toggle_sidebar)
        self.btn_show_sidebar.setVisible(False)
        hl.addWidget(self.btn_show_sidebar)

        cfg = self.cfg_getter()
        self.lbl_model = QLabel(f"⚡ {cfg.get('lm_model','—').split('/')[-1]}")
        self.lbl_model.setStyleSheet("color: #888888; font-size: 12px; font-family: 'Segoe UI', Arial;")
        hl.addStretch()
        hl.addWidget(self.lbl_model)
        hl.addSpacing(8)

        self.cmb_tone = QComboBox()
        self.cmb_tone.addItems(list(ChatWorker.TONES.keys()))
        self.cmb_tone.setFixedHeight(30)
        hl.addWidget(self.cmb_tone)

        cl.addWidget(header)

        # Area messaggi con scroll
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        self.msg_container = QWidget()
        self.msg_container.setObjectName("msg_container")
        self.msg_layout = QVBoxLayout(self.msg_container)
        self.msg_layout.setSpacing(0)
        self.msg_layout.setContentsMargins(0, 16, 0, 16)
        self.msg_layout.addStretch()

        self.scroll.setWidget(self.msg_container)
        cl.addWidget(self.scroll, 1)

        # Input area
        input_area = QFrame()
        input_area.setStyleSheet("background-color: #212121; border-top: 1px solid #2A2A2A;")
        ia = QVBoxLayout(input_area)
        ia.setContentsMargins(20, 12, 20, 14)
        ia.setSpacing(8)

        # Input + send
        input_row = QHBoxLayout()
        input_row.setSpacing(10)
        self.inp = QTextEdit()
        self.inp.setObjectName("msg_input")
        self.inp.setPlaceholderText("Scrivi un messaggio... (Invio per inviare, Shift+Invio per andare a capo)")
        self.inp.setFixedHeight(52)
        self.inp.installEventFilter(self)
        input_row.addWidget(self.inp)
        self.btn_send = QPushButton("↑")
        self.btn_send.setObjectName("send_btn")
        self.btn_send.setFixedSize(44, 44)
        self.btn_send.clicked.connect(self._send)
        input_row.addWidget(self.btn_send, 0, Qt.AlignmentFlag.AlignBottom)
        ia.addLayout(input_row)

        hint = QLabel("L'assistente può commettere errori. Verificare le informazioni importanti.")
        hint.setStyleSheet("color: #444444; font-size: 11px; font-family: 'Segoe UI', Arial;")
        hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        ia.addWidget(hint)

        cl.addWidget(input_area)
        root.addWidget(chat_widget, 1)

        # Messaggio benvenuto
        self._add_bubble("assistant", self.WELCOME)

    # ── Sidebar ───────────────────────────────────────────
    def _toggle_sidebar(self):
        self.sidebar_visible = not self.sidebar_visible
        self.sidebar.setVisible(self.sidebar_visible)
        self.btn_show_sidebar.setVisible(not self.sidebar_visible)
        self.btn_toggle.setText("◀" if self.sidebar_visible else "▶")

    def _new_chat(self):
        """Salva la chat corrente e ne inizia una nuova."""
        self._save_current_chat()
        self.history = []
        self.lbl_attach.setText("")
        self._clear_bubbles()
        self._add_bubble("assistant", self.WELCOME)

    def _save_current_chat(self):
        """Salva la chat corrente nella cronologia se ha messaggi visibili e non è già salvata."""
        visible = [m for m in self.history
                   if m["role"] == "user" and not m["content"].startswith("[CONTESTO")]
        if not visible:
            return
        # Evita duplicati — controlla se la stessa chat è già in cima
        title = next((m["content"][:35] + ("…" if len(m["content"]) > 35 else "")
                      for m in self.history if m["role"] == "user"
                      and not m["content"].startswith("[CONTESTO")), "Chat")
        if self.all_chats and self.all_chats[0]["title"] == title:
            # Aggiorna la chat esistente
            self.all_chats[0]["history"] = list(self.history)
        else:
            self.all_chats.insert(0, {"title": title, "history": list(self.history)})
        self._refresh_history_list()

    def _refresh_history_list(self):
        self.list_history.clear()
        for i, chat in enumerate(self.all_chats):
            item = QListWidgetItem()
            item.setData(Qt.ItemDataRole.UserRole, i)
            item.setSizeHint(QSize(220, 44))
            item.setText(f"  💬  {chat['title']}")
            self.list_history.addItem(item)

    def _show_context_menu(self, pos):
        from PyQt6.QtWidgets import QMenu
        item = self.list_history.itemAt(pos)
        if not item:
            return
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background-color: #2F2F2F;
                border: 1px solid #3F3F3F;
                color: #ECECEC;
                font-family: 'Segoe UI', Arial;
                font-size: 13px;
                padding: 4px;
            }
            QMenu::item { padding: 8px 20px; border-radius: 4px; }
            QMenu::item:selected { background-color: #EF4444; color: white; }
        """)
        act_del = menu.addAction("🗑  Elimina questa chat")
        act_load = menu.addAction("↩  Carica questa chat")
        action = menu.exec(self.list_history.mapToGlobal(pos))
        idx = self.list_history.row(item)
        if action == act_del:
            self._delete_chat(idx)
        elif action == act_load:
            self._load_chat(item)

    def _delete_selected_chat(self):
        row = self.list_history.currentRow()
        if row >= 0:
            self._delete_chat(row)
        else:
            self._add_bubble("assistant", "Seleziona prima una chat dalla cronologia per eliminarla.")

    def _delete_chat(self, idx):
        if 0 <= idx < len(self.all_chats):
            self.all_chats.pop(idx)
            self._refresh_history_list()

    def _load_chat(self, item):
        """Carica una chat dalla cronologia SENZA salvare quella corrente."""
        idx = self.list_history.row(item)
        if 0 <= idx < len(self.all_chats):
            self.history = list(self.all_chats[idx]["history"])
            self.lbl_attach.setText("")
            self._clear_bubbles()
            self._add_bubble("assistant", self.WELCOME)
            for msg in self.history:
                if msg["role"] == "user" and not msg["content"].startswith("[CONTESTO"):
                    self._add_bubble("user", msg["content"])
                elif msg["role"] == "assistant":
                    self._add_bubble("assistant", msg["content"])
            self._add_bubble("assistant", "Cronologia ripristinata. Può continuare la conversazione.")

    def _clear_bubbles(self):
        while self.msg_layout.count() > 1:
            item = self.msg_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

    def _clear_chat(self):
        self._save_current_chat()
        self.history = []
        self.lbl_attach.setText("")
        self._clear_bubbles()
        self._add_bubble("assistant", "Chat cancellata. Come posso aiutarla?")

    # ── Messaggi ──────────────────────────────────────────
    def _add_bubble(self, role, text):
        wrapper = QWidget()
        wrapper.setStyleSheet("background-color: transparent;")
        wl = QHBoxLayout(wrapper)
        wl.setContentsMargins(0, 4, 0, 4)

        if role == "user":
            wl.addStretch()
            bubble = QLabel(text)
            bubble.setWordWrap(True)
            bubble.setMaximumWidth(600)
            bubble.setMinimumWidth(120)
            bubble.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            bubble.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
            bubble.setStyleSheet("""
                QLabel {
                    background-color: #2F2F2F;
                    color: #ECECEC;
                    border-radius: 18px;
                    padding: 12px 16px;
                    font-family: 'Segoe UI', Arial, sans-serif;
                    font-size: 14px;
                }
            """)
            wl.addWidget(bubble)
            wl.setContentsMargins(80, 4, 24, 4)
        else:
            avatar = QLabel("✦")
            avatar.setFixedSize(32, 32)
            avatar.setAlignment(Qt.AlignmentFlag.AlignCenter)
            avatar.setStyleSheet("""
                QLabel {
                    background-color: #19C37D;
                    color: white;
                    border-radius: 16px;
                    font-size: 14px;
                    font-weight: bold;
                }
            """)
            wl.addWidget(avatar, 0, Qt.AlignmentFlag.AlignTop)
            wl.addSpacing(10)

            msg_col = QVBoxLayout()
            msg_col.setSpacing(4)
            name_lbl = QLabel("Assistente")
            name_lbl.setStyleSheet("color: #888888; font-size: 11px; font-family: 'Segoe UI', Arial;")
            msg_col.addWidget(name_lbl)

            bubble = QLabel(text)
            bubble.setWordWrap(True)
            bubble.setMaximumWidth(650)
            bubble.setMinimumWidth(300)
            bubble.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
            bubble.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
            bubble.setStyleSheet("""
                QLabel {
                    color: #ECECEC;
                    background-color: transparent;
                    font-family: 'Segoe UI', Arial, sans-serif;
                    font-size: 14px;
                    padding: 2px 0;
                }
            """)
            msg_col.addWidget(bubble)
            wl.addLayout(msg_col)
            wl.addStretch()
            wl.setContentsMargins(20, 4, 80, 4)

        self.msg_layout.insertWidget(self.msg_layout.count() - 1, wrapper)
        QApplication.processEvents()
        self.scroll.verticalScrollBar().setValue(self.scroll.verticalScrollBar().maximum())

    # ── Allegato ──────────────────────────────────────────
    def _attach_transcript(self):
        if self.get_transcript:
            transcript = self.get_transcript()
            if transcript and transcript.strip():
                self.history.append({"role":"user","content":f"[CONTESTO — Trascrizione]:\n{transcript}\n[Fine]"})
                self.history.append({"role":"assistant","content":"Ho acquisito la trascrizione."})
                self.lbl_attach.setText("✅ Allegata")
                self._add_bubble("assistant",
                    "Ho acquisito la trascrizione della riunione corrente. "
                    "Posso analizzarla, riassumerla, estrarre decisioni o action items.")
            else:
                self.lbl_attach.setText("⚠️ Nessuna trascrizione")

    # ── Invio ─────────────────────────────────────────────
    def eventFilter(self, obj, event):
        from PyQt6.QtCore import QEvent
        if obj == self.inp and event.type() == QEvent.Type.KeyPress:
            if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                if event.modifiers() == Qt.KeyboardModifier.ShiftModifier:
                    return False
                self._send()
                return True
        return super().eventFilter(obj, event)

    def _send(self):
        text = self.inp.toPlainText().strip()
        if not text or (self.worker and self.worker.isRunning()):
            return
        self.inp.clear()
        self.btn_send.setEnabled(False)
        self.history.append({"role":"user","content":text})
        self._add_bubble("user", text)

        # Indicatore typing
        self._typing_widget = QWidget()
        self._typing_widget.setStyleSheet("background-color: transparent;")
        tl = QHBoxLayout(self._typing_widget)
        tl.setContentsMargins(20, 4, 80, 4)
        av = QLabel("✦")
        av.setFixedSize(32,32)
        av.setAlignment(Qt.AlignmentFlag.AlignCenter)
        av.setStyleSheet("background-color:#19C37D;color:white;border-radius:16px;font-size:14px;")
        tl.addWidget(av, 0, Qt.AlignmentFlag.AlignTop)
        tl.addSpacing(10)
        dots = QLabel("● ● ●")
        dots.setStyleSheet("color:#444444;font-size:18px;letter-spacing:4px;")
        tl.addWidget(dots)
        tl.addStretch()
        self.msg_layout.insertWidget(self.msg_layout.count()-1, self._typing_widget)
        QApplication.processEvents()
        self.scroll.verticalScrollBar().setValue(self.scroll.verticalScrollBar().maximum())

        cfg = self.cfg_getter()
        self.lbl_model.setText(f"⚡ {cfg.get('lm_model','—').split('/')[-1]}")
        self.worker = ChatWorker(self.history[-20:], cfg, self.cmb_tone.currentText())
        self.worker.token.connect(self._on_response)
        self.worker.finished.connect(self._on_done)
        self.worker.error.connect(self._on_error)
        self.worker.start()

    def _on_response(self, text):
        if hasattr(self, '_typing_widget'):
            self._typing_widget.setParent(None)
            del self._typing_widget
        self.history.append({"role":"assistant","content":text})
        self._add_bubble("assistant", text)

    def _on_done(self):
        self.btn_send.setEnabled(True)
        self.inp.setFocus()

    def _on_error(self, err):
        if hasattr(self, '_typing_widget'):
            self._typing_widget.setParent(None)
            del self._typing_widget
        self.history.append({"role":"assistant","content":f"Errore: {err}"})
        self._add_bubble("assistant", f"⚠️ {err}")
        self._on_done()

    def closeEvent(self, event):
        """Salva la chat corrente alla chiusura."""
        self._save_current_chat()
        super().closeEvent(event)


# ──────────────────────────────────────────────────────
# FINESTRA PRINCIPALE
# ──────────────────────────────────────────────────────
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Meeting Recorder")
        self.setMinimumSize(900, 680)
        self.resize(1100, 750)

        self._build_ui()
        self._connect_signals()

    def _build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # ── Top bar ──────────────────────────────────
        topbar = QFrame()
        topbar.setFixedHeight(64)
        topbar.setStyleSheet("background-color: #0D1117; border-bottom: 1px solid #1E293B;")
        tbl = QHBoxLayout(topbar)
        tbl.setContentsMargins(24, 0, 24, 0)

        dot_red   = QLabel("●"); dot_red.setStyleSheet("color:#EF4444; font-size:14px;")
        dot_yel   = QLabel("●"); dot_yel.setStyleSheet("color:#F59E0B; font-size:14px;")
        dot_grn   = QLabel("●"); dot_grn.setStyleSheet("color:#10B981; font-size:14px;")
        for d in [dot_red, dot_yel, dot_grn]:
            tbl.addWidget(d)
        tbl.addSpacing(16)

        logo = QLabel("🎙  Meeting Recorder")
        logo.setFont(QFont("SF Pro Display", 16, QFont.Weight.Bold))
        logo.setStyleSheet("color: #F1F5F9;")
        tbl.addWidget(logo)
        tbl.addStretch()

        # Pulsante chat
        self.btn_chat = QPushButton("💬  Chat AI")
        self.btn_chat.setFixedHeight(36)
        self.btn_chat.setFixedWidth(110)
        self.btn_chat.setStyleSheet("""
            QPushButton {
                background-color: #6D28D9;
                color: white;
                border: none;
                border-radius: 8px;
                font-size: 13px;
                font-weight: 600;
                font-family: Arial;
            }
            QPushButton:hover { background-color: #5B21B6; }
            QPushButton:pressed { background-color: #4C1D95; }
        """)
        self.btn_chat.clicked.connect(self._open_chat)
        tbl.addWidget(self.btn_chat)
        tbl.addSpacing(16)

        dev = QLabel("Sviluppato da  Marco Bonometti")
        dev.setStyleSheet("color: #475569; font-size: 12px;")
        tbl.addWidget(dev)

        root.addWidget(topbar)

        # ── Tabs ──────────────────────────────────────
        self.settings_tab = SettingsTab()

        self.tabs = QTabWidget()
        self.tabs.setTabPosition(QTabWidget.TabPosition.North)

        self.record_tab  = RecordTab(self.settings_tab.get_config)
        self.reports_tab = ReportsTab(self.settings_tab.get_config)

        self.tabs.addTab(self.record_tab,  "⏺   Registra")
        self.tabs.addTab(self.reports_tab, "📂  Report")
        self.tabs.addTab(self.settings_tab,"⚙️   Impostazioni")

        root.addWidget(self.tabs)

        # ── Status bar ────────────────────────────────
        self.status = QStatusBar()
        self.setStatusBar(self.status)
        self.status.showMessage(
            f"  Meeting Recorder  ·  Sviluppato da Marco Bonometti  ·  "
            f"Pronto  ·  {datetime.date.today().strftime('%d/%m/%Y')}"
        )

    def _connect_signals(self):
        self.record_tab.session_done.connect(self.reports_tab.add_session)
        self.settings_tab.config_changed.connect(self._on_config_changed)

    def _open_chat(self):
        """Apre la finestra chat passando la trascrizione corrente se disponibile."""
        def get_transcript():
            return self.record_tab.txt_transcript.toPlainText()
        chat = ChatWindow(
            cfg_getter=self.settings_tab.get_config,
            get_transcript=get_transcript,
            parent=self,
        )
        chat.exec()

    def _on_config_changed(self, cfg):
        self.status.showMessage(f"  Impostazioni aggiornate — Modello: {cfg['lm_model']}")


# ──────────────────────────────────────────────────────
# ENTRY POINT
# ──────────────────────────────────────────────────────
def main():
    log_path = Path.home() / "meeting_recorder_log.txt"
    import traceback

    # Cattura TUTTI i crash inclusi quelli nei thread Qt
    def global_exception_hook(exctype, value, tb):
        err = "".join(traceback.format_exception(exctype, value, tb))
        try:
            log_path.write_text(err, encoding="utf-8")
        except Exception:
            pass
        try:
            from PyQt6.QtWidgets import QMessageBox
            _app = QApplication.instance()
            if _app:
                msg = QMessageBox()
                msg.setIcon(QMessageBox.Icon.Critical)
                msg.setWindowTitle("Errore")
                msg.setText(f"Errore:\n{str(value)}\n\nLog: {log_path}")
                msg.exec()
        except Exception:
            pass

    sys.excepthook = global_exception_hook

    # Cattura crash nei QThread
    import threading
    def thread_exception_hook(args):
        err = "".join(traceback.format_exception(args.exc_type, args.exc_value, args.exc_traceback))
        try:
            log_path.write_text(err, encoding="utf-8")
        except Exception:
            pass
    threading.excepthook = thread_exception_hook

    try:
        app = QApplication(sys.argv)
        app.setApplicationName("Meeting Recorder")
        app.setOrganizationName("Marco Bonometti")
        app.setStyle("Fusion")
        app.setStyleSheet(STYLE)

        pal = QPalette()
        pal.setColor(QPalette.ColorRole.Window,          QColor("#0F1117"))
        pal.setColor(QPalette.ColorRole.WindowText,      QColor("#E2E8F0"))
        pal.setColor(QPalette.ColorRole.Base,            QColor("#161B27"))
        pal.setColor(QPalette.ColorRole.AlternateBase,   QColor("#1E293B"))
        pal.setColor(QPalette.ColorRole.Text,            QColor("#E2E8F0"))
        pal.setColor(QPalette.ColorRole.ButtonText,      QColor("#E2E8F0"))
        pal.setColor(QPalette.ColorRole.Highlight,       QColor("#2563EB"))
        pal.setColor(QPalette.ColorRole.HighlightedText, QColor("#FFFFFF"))
        app.setPalette(pal)

        win = MainWindow()
        win.show()

        sys.exit(app.exec())

    except Exception as e:
        err_msg = traceback.format_exc()
        try:
            log_path.write_text(err_msg, encoding="utf-8")
        except Exception:
            pass
        try:
            _app = QApplication.instance() or QApplication(sys.argv)
            from PyQt6.QtWidgets import QMessageBox
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.setWindowTitle("Errore avvio")
            msg.setText(
                f"Errore:\n{str(e)}\n\n"
                f"Log salvato in:\n{log_path}"
            )
            msg.exec()
        except Exception:
            pass
        sys.exit(1)


if __name__ == "__main__":
    main()
